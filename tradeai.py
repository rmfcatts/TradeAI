import os
import sys
import openpyxl
import docx
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.document_loaders import TextLoader
from langchain_community.vectorstores import FAISS
from langchain.chains import RetrievalQA
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document  # Import the Document class
from constants import APIKEY, folder_path

# Ensure the OpenAI API key is set
openai_api_key = os.getenv("OPENAI_API_KEY", APIKEY)
if openai_api_key == "sk-your-api-key-here":
    print("Please set your OpenAI API key in the environment variable or directly in the script.")
    os.sys.exit(1)

# Memory file path
memory_file = os.path.join(folder_path, 'user_memory.txt')

# Function to read Excel files using openpyxl
def read_excel(file_path):
    text = ""
    workbook = openpyxl.load_workbook(file_path, read_only=True)
    for sheet in workbook:
        for row in sheet.iter_rows(values_only=True):
            text += " ".join([str(cell) for cell in row if cell is not None]) + "\n"
    return text

# Function to read DOCX files
def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([paragraph.text for paragraph in doc.paragraphs])

# Function to read CSV files using pandas
def read_csv(file_path):
    import pandas as pd  # Import pandas here
    df = pd.read_csv(file_path)
    return df.to_string(index=False)

# Function to read the memory file
def read_memory_file(file_path):
    if os.path.exists(file_path):
        with open(file_path, 'r') as f:
            return f.read()
    return ""

# Function to update the memory file
def update_memory_file(file_path, question, answer):
    with open(file_path, 'a') as f:
        f.write(f"Q: {question}\nA: {answer}\n\n")

# Function to load documents from the folder
def load_documents_from_folder(folder_path):
    documents = []
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        if os.path.isfile(file_path):
            if filename.endswith('.xlsx'):
                text = read_excel(file_path)
                documents.append(Document(page_content=text, metadata={"source": filename}))
            elif filename.endswith('.docx'):
                text = read_docx(file_path)
                documents.append(Document(page_content=text, metadata={"source": filename}))
            elif filename.endswith('.csv'):
                text = read_csv(file_path)
                documents.append(Document(page_content=text, metadata={"source": filename}))
            elif filename.endswith('.txt'):
                loader = TextLoader(file_path)
                text_documents = loader.load()
                for text_doc in text_documents:
                    documents.append(Document(page_content=text_doc.page_content, metadata={"source": filename}))
            else:
                continue
    return documents

# Load and split documents
documents = load_documents_from_folder(folder_path)

# Load user memory file content and add to documents
user_memory_content = read_memory_file(memory_file)
if user_memory_content:
    documents.append(Document(page_content=user_memory_content, metadata={"source": "user_memory.txt"}))

text_splitter = RecursiveCharacterTextSplitter(chunk_size=1000, chunk_overlap=200)
docs = text_splitter.split_documents(documents)

# Index documents
embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)
vectorstore = FAISS.from_documents(docs, embeddings)

# Create RetrievalQA chain
qa_chain = RetrievalQA.from_chain_type(
    llm=ChatOpenAI(model="gpt-4", openai_api_key=openai_api_key),
    retriever=vectorstore.as_retriever(),
    return_source_documents=True
)

def query_documents(question):
    result = qa_chain.invoke({"query": question})
    if result['source_documents']:
        source = result['source_documents'][0].metadata["source"]
        return result['result'], source
    return None, None

# Ensure the user's question is provided
if len(sys.argv) > 1:
    user_question = sys.argv[1].strip().lower()
else:
    print("Please provide a question as a command line argument.")
    sys.exit(1)

# Query documents
answer, source = query_documents(user_question)

if not answer or "I don't know" in answer:
    print("I'm sorry, but I don't have that information.")
else:
    print(f"{answer}")

    # Store new information in memory
    update_memory_file(memory_file, user_question, answer)
